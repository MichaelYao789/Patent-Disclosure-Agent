"""
将 README.md 转换为可编辑的 Word 文档
=====================================
运行：python generate_readme_docx.py
输出：output/README.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt

def markdown_to_docx(md_path: str, docx_path: str):
    """
    将 Markdown 文件转换为 Word 文档
    
    参数：
        md_path: Markdown 文件路径
        docx_path: 输出 Word 文件路径
    """
    if not os.path.exists(md_path):
        print(f"未找到 {md_path}")
        return
    
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()
    
    doc = Document()
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3
    
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # 标题处理
        if line.startswith('###### '):
            doc.add_heading(line[7:], level=6)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        # 代码块
        elif line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        # 列表
        elif re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            doc.add_paragraph(text, style='List Number')
        else:
            doc.add_paragraph(line)
        i += 1
    
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"README.docx 已生成：{docx_path}")

if __name__ == "__main__":
    markdown_to_docx("README.md", "output/README.docx")
