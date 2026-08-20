"""
Word 文档生成器
===============
将 Markdown 内容转换为 Word 格式，包含交底书正文、检索资料和可编辑源码附录。

核心类：
- DocxGenerator：Word 生成器
"""

import os
import re
import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class DocxGenerator:
    """
    Word 文档生成器
    
    使用 python-docx 将 Markdown 内容转换为标准 Word 文档。
    """
    
    def __init__(self, output_dir: str):
        """
        初始化
        
        参数：
            output_dir: 输出目录
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, state: Dict[str, Any], markdown_text: str) -> str:
        """
        生成 Word 文档
        
        参数：
            state: 执行状态（含检索结果）
            markdown_text: 交底书 Markdown 文本
        
        返回：
            str: Word 文件路径
        """
        doc = Document()
        self._set_default_style(doc)
        
        # 封面
        doc.add_heading('发明专利交底书', level=0)
        p = doc.add_paragraph('（由 AI 辅助生成，最终请专利代理人审核）')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 正文
        self._add_markdown_lines(doc, markdown_text)
        
        # 检索资料
        self._add_retrieval_results(doc, state.get("retrieval_results", {}))
        
        # 可编辑源码附录
        self._add_figures_appendix(doc)
        
        # 保存
        path = os.path.join(self.output_dir, "专利交底书.docx")
        doc.save(path)
        logger.info(f"Word 文档已保存: {path}")
        return path

    def _set_default_style(self, doc):
        """设置默认字体和段落格式"""
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    def _add_markdown_lines(self, doc, md_text: str):
        """将 Markdown 逐行解析写入 Word"""
        lines = md_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line.strip():
                i += 1
                continue
            # 标题
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # 列表
            elif re.match(r'^\s*[-*]\s+', line):
                text = re.sub(r'^\s*[-*]\s+', '', line)
                doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\s*\d+\.\s+', line):
                text = re.sub(r'^\s*\d+\.\s+', '', line)
                doc.add_paragraph(text, style='List Number')
            # 代码块
            elif line.strip().startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            else:
                doc.add_paragraph(line)
            i += 1

    def _add_retrieval_results(self, doc, retrieval_results: Dict[str, Any]):
        """添加检索资料汇总"""
        if not retrieval_results:
            return
        doc.add_heading('附：检索资料汇总（真实来源）', level=1)
        for step_id, data in retrieval_results.items():
            doc.add_heading(f"检索步骤 {step_id}: {data['tool']} / {data['query']}", level=2)
            for item in data.get("results", []):
                p = doc.add_paragraph()
                p.add_run(f"{item.get('title', '')}\n").bold = True
                p.add_run(f"{item.get('abstract', '')[:400]}\n")
                p.add_run(f"来源：{item.get('source', '')}")

    def _add_figures_appendix(self, doc):
        """添加可编辑流程图源码"""
        dot_files = [f for f in os.listdir(self.output_dir) if f.endswith('.dot')]
        if dot_files:
            doc.add_heading('附录：可编辑流程图源码（DOT 格式）', level=1)
            for fname in dot_files:
                doc.add_heading(fname, level=2)
                with open(os.path.join(self.output_dir, fname), 'r', encoding='utf-8') as f:
                    code = f.read()
                p = doc.add_paragraph()
                run = p.add_run(code)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
